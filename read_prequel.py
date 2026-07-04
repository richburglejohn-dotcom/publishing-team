from docx import Document

doc = Document("Copy of Culinary Crescendo - The Scarlet Foundation (Prequel) v10.docx")
output = []
for para in doc.paragraphs[200:]:
    if para.text.strip():
        output.append(para.text)

with open("outputs/prequel_text_extract_part2.txt", "w", encoding="utf-8") as f:
    f.write("\n\n".join(output))

print(f"Done! Extracted {len(output)} more paragraphs.")
print(f"Total paragraphs in doc: {len(doc.paragraphs)}")
